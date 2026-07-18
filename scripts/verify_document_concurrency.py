#!/usr/bin/env python3
"""Fail-closed, auditable document upload and concurrency verification."""

from __future__ import annotations

import argparse
import asyncio
from datetime import datetime, timezone
import json
import math
import os
from pathlib import Path
import sys
import tempfile
import time
from typing import Any
import uuid

import httpx


ROOT = Path(__file__).resolve().parents[1]
BACKEND = ROOT / "backend"
if str(BACKEND) not in sys.path:
    sys.path.insert(0, str(BACKEND))

FORMATS = ("pdf", "docx", "xlsx", "txt")
HEALTH_SAMPLES_PER_FORMAT = 20
PROJECT_SAMPLES_PER_FORMAT = 10
DEFAULT_JSON_OUTPUT = ROOT / "artifacts" / "document_concurrency_raw.json"
DEFAULT_MARKDOWN_OUTPUT = ROOT / "artifacts" / "document_concurrency_summary.md"


def _utc_now() -> str:
    return datetime.now(timezone.utc).isoformat().replace("+00:00", "Z")


def _write_pdf(path: Path, scale: int) -> None:
    from pypdf import PdfWriter
    from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

    writer = PdfWriter()
    for page_index in range(max(2, scale * 10)):
        page = writer.add_blank_page(width=612, height=792)
        font = DictionaryObject(
            {
                NameObject("/Type"): NameObject("/Font"),
                NameObject("/Subtype"): NameObject("/Type1"),
                NameObject("/BaseFont"): NameObject("/Helvetica"),
            }
        )
        page[NameObject("/Resources")] = DictionaryObject(
            {
                NameObject("/Font"): DictionaryObject(
                    {NameObject("/F1"): writer._add_object(font)}
                )
            }
        )
        stream = DecodedStreamObject()
        lines = [
            "BT /F1 10 Tf 54 740 Td",
            *[
                f"(Document verification page {page_index + 1} line {line + 1}) Tj 0 -12 Td"
                for line in range(35)
            ],
            "ET",
        ]
        stream.set_data("\n".join(lines).encode("ascii"))
        page[NameObject("/Contents")] = writer._add_object(stream)
    with path.open("wb") as output:
        writer.write(output)


def _write_docx(path: Path, scale: int) -> None:
    from docx import Document

    document = Document()
    document.add_heading("并发验证文档", level=1)
    for index in range(scale * 1000):
        document.add_paragraph(f"第 {index + 1} 段：真实 DOCX 解析与并发探针内容。")
    table = document.add_table(rows=10, cols=3)
    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            cell.text = f"表格-{row_index}-{column_index}"
    document.save(path)


def _write_xlsx(path: Path, scale: int) -> None:
    from openpyxl import Workbook

    workbook = Workbook(write_only=False)
    first = workbook.active
    first.title = "主数据"
    second = workbook.create_sheet("补充数据")
    for index in range(scale * 1000):
        first.append([f"项目-{index}", index, None, f"内容-{index}"])
        second.append([f"补充-{index}", index / 10, None])
    workbook.save(path)


def generate_samples(directory: Path, scale: int = 3) -> dict[str, Path]:
    """Generate large, structurally valid samples with official libraries."""

    if scale < 1:
        raise ValueError("scale must be at least 1")
    directory.mkdir(parents=True, exist_ok=True)
    samples = {
        "pdf": directory / "document-concurrency.pdf",
        "docx": directory / "document-concurrency.docx",
        "xlsx": directory / "document-concurrency.xlsx",
        "txt": directory / "document-concurrency.txt",
    }
    _write_pdf(samples["pdf"], scale)
    _write_docx(samples["docx"], scale)
    _write_xlsx(samples["xlsx"], scale)
    samples["txt"].write_text(
        "".join(f"第 {index + 1} 行：真实 TXT 中文并发解析样本。\n" for index in range(scale * 10000)),
        encoding="utf-8",
    )
    return samples


def validate_generated_samples(samples: dict[str, Path]) -> dict[str, dict[str, Any]]:
    """Reopen samples with official libraries and verify extracted content."""

    from app.tools.document_extractors import extract_document

    validation: dict[str, dict[str, Any]] = {}
    from pypdf import PdfReader
    from docx import Document
    from openpyxl import load_workbook

    pdf = PdfReader(samples["pdf"])
    validation["pdf"] = {"official_reopen": True, "page_count": len(pdf.pages)}
    docx = Document(samples["docx"])
    validation["docx"] = {
        "official_reopen": True,
        "paragraph_count": len(docx.paragraphs),
        "table_count": len(docx.tables),
    }
    workbook = load_workbook(samples["xlsx"], read_only=True, data_only=True)
    try:
        validation["xlsx"] = {
            "official_reopen": True,
            "sheet_count": len(workbook.sheetnames),
        }
    finally:
        workbook.close()
    with samples["txt"].open("r", encoding="utf-8") as stream:
        stream.read(1)
    validation["txt"] = {"official_reopen": True, "encoding": "utf-8"}

    for kind, path in samples.items():
        result = extract_document(str(path))
        if "error" in result:
            raise RuntimeError(f"generated {kind} sample failed extraction: {result['error']}")
        validation[kind]["extracted_char_count"] = len(result["content"])
        validation[kind]["size_bytes"] = path.stat().st_size
    return validation


def _nearest_rank(values: list[float], percentile: float) -> float:
    if not values:
        return 0.0
    ordered = sorted(values)
    index = max(0, math.ceil(percentile * len(ordered)) - 1)
    return ordered[index]


def _probe_summary(samples: list[dict[str, Any]]) -> dict[str, Any]:
    durations = [float(item["duration_seconds"]) for item in samples]
    failures = sum(
        1
        for item in samples
        if item.get("error") or not (200 <= int(item.get("status_code") or 0) < 300)
    )
    ordered = sorted(durations)
    if not ordered:
        p50 = 0.0
    elif len(ordered) % 2:
        p50 = ordered[len(ordered) // 2]
    else:
        middle = len(ordered) // 2
        p50 = (ordered[middle - 1] + ordered[middle]) / 2
    return {
        "samples": len(samples),
        "failures": failures,
        "p50_seconds": round(p50, 6),
        "p95_seconds": round(_nearest_rank(durations, 0.95), 6),
        "max_seconds": round(max(durations, default=0.0), 6),
    }


def build_summary(report: dict[str, Any]) -> dict[str, Any]:
    return {
        "health": _probe_summary(report.get("probes", {}).get("health", [])),
        "projects": _probe_summary(report.get("probes", {}).get("projects", [])),
        "total_duration_seconds": report.get("total_duration_seconds", 0.0),
        "formats": report.get("formats", {}),
    }


def render_markdown(report: dict[str, Any], summary: dict[str, Any]) -> str:
    lines = [
        "# 文档解析并发验证摘要",
        "",
        f"- 开始：{report.get('started_at', '')}",
        f"- 结束：{report.get('finished_at', '')}",
        f"- 总耗时：{float(summary['total_duration_seconds']):.3f}s",
        "",
        "| 探针 | 样本 | 失败 | p50(s) | p95(s) | max(s) |",
        "|---|---:|---:|---:|---:|---:|",
    ]
    for name in ("health", "projects"):
        item = summary[name]
        lines.append(
            f"| {name} | {item['samples']} | {item['failures']} | "
            f"{item['p50_seconds']:.3f} | {item['p95_seconds']:.3f} | {item['max_seconds']:.3f} |"
        )
    lines.extend(
        [
            "",
            "| 格式 | 结果 | 知识块 | 上传/解析耗时(s) |",
            "|---|---|---:|---:|",
        ]
    )
    for kind, item in summary["formats"].items():
        lines.append(
            f"| {kind} | {item.get('status')} | {item.get('document_count', 0)} | "
            f"{float(item.get('duration_seconds', 0)):.3f} |"
        )
    if report.get("errors"):
        lines.extend(["", "## 失败", "", *[f"- {error}" for error in report["errors"]]])
    return "\n".join(lines) + "\n"


def validate_report(report: dict[str, Any]) -> list[str]:
    errors: list[str] = []
    formats = report.get("formats", {})
    for kind in FORMATS:
        item = formats.get(kind)
        if not item:
            errors.append(f"{kind}: missing format result")
            continue
        if item.get("status") != "completed":
            errors.append(f"{kind}: status {item.get('status')} is not completed")
        if int(item.get("document_count") or 0) <= 0:
            errors.append(f"{kind}: document_count is empty")
        if item.get("cleanup_status") != "completed":
            errors.append(f"{kind}: cleanup failed")

    thresholds = {
        "health": (HEALTH_SAMPLES_PER_FORMAT, 1.0, 2.0),
        "projects": (PROJECT_SAMPLES_PER_FORMAT, 2.0, 3.0),
    }
    for probe_name, (required_per_format, p95_limit, max_limit) in thresholds.items():
        samples = report.get("probes", {}).get(probe_name, [])
        for kind in FORMATS:
            count = sum(1 for sample in samples if sample.get("format") == kind)
            if count < required_per_format:
                errors.append(
                    f"{probe_name}/{kind}: requires {required_per_format} samples, got {count}"
                )
        summary = _probe_summary(samples)
        if summary["failures"]:
            errors.append(f"{probe_name}: {summary['failures']} non-2xx/error samples")
        if summary["p95_seconds"] >= p95_limit:
            errors.append(
                f"{probe_name}: p95 {summary['p95_seconds']:.3f}s exceeds {p95_limit:.3f}s"
            )
        if summary["max_seconds"] >= max_limit:
            errors.append(
                f"{probe_name}: max {summary['max_seconds']:.3f}s exceeds {max_limit:.3f}s"
            )
    return errors


async def _probe(
    client: httpx.AsyncClient, path: str, kind: str, name: str
) -> dict[str, Any]:
    started_at = _utc_now()
    started = time.perf_counter()
    status_code = None
    error = None
    try:
        response = await client.get(path)
        status_code = response.status_code
        if not 200 <= response.status_code < 300:
            error = response.text[:500]
    except Exception as exc:
        error = f"{type(exc).__name__}: {exc}"
    return {
        "probe": name,
        "format": kind,
        "started_at": started_at,
        "duration_seconds": round(time.perf_counter() - started, 6),
        "status_code": status_code,
        "error": error,
    }


async def _poll_completion(client: httpx.AsyncClient, kb_id: int, timeout: float) -> dict[str, Any]:
    deadline = time.monotonic() + timeout
    while time.monotonic() < deadline:
        progress = await client.get(f"/api/v1/knowledge/{kb_id}/progress")
        if not 200 <= progress.status_code < 300:
            raise RuntimeError(f"progress non-2xx: {progress.status_code} {progress.text[:500]}")
        data = progress.json().get("data") or {}
        if data.get("status") == "failed" or data.get("error"):
            raise RuntimeError(data.get("error") or "knowledge processing failed")
        if data.get("status") == "completed":
            detail = await client.get(f"/api/v1/knowledge/{kb_id}")
            if not 200 <= detail.status_code < 300:
                raise RuntimeError(f"detail non-2xx: {detail.status_code} {detail.text[:500]}")
            record = detail.json().get("data") or {}
            if int(record.get("document_count") or 0) <= 0:
                raise RuntimeError("completed upload has empty document_count")
            return record
        await asyncio.sleep(0.1)
    raise TimeoutError(f"knowledge upload {kb_id} timed out after {timeout}s")


async def _list_owned_records(
    client: httpx.AsyncClient, owned_name: str
) -> list[dict[str, Any]]:
    response = await client.get("/api/v1/knowledge")
    if not 200 <= response.status_code < 300:
        raise RuntimeError(f"cleanup list non-2xx: {response.status_code} {response.text[:500]}")
    data = response.json().get("data") or []
    if isinstance(data, dict):
        data = data.get("items") or []
    return [record for record in data if record.get("name") == owned_name]


async def _cleanup_owned_records(
    client: httpx.AsyncClient, owned_name: str, known_id: int | None
) -> list[int]:
    """Find by unique name so response timeouts cannot orphan accepted uploads."""

    records = await _list_owned_records(client, owned_name)
    owned_ids = {int(record["id"]) for record in records}
    if known_id is not None:
        owned_ids.add(int(known_id))
    statuses = {int(record["id"]): record.get("status") for record in records}
    for kb_id in sorted(owned_ids):
        if statuses.get(kb_id) == "processing":
            stop = await client.post(f"/api/v1/knowledge/{kb_id}/stop")
            # A terminal-state race may make stop return 400; deletion still
            # remains authoritative and is verified below.
            if stop.status_code not in {200, 400}:
                raise RuntimeError(
                    f"cleanup stop non-2xx: {stop.status_code} {stop.text[:500]}"
                )
        cleanup = await client.delete(f"/api/v1/knowledge/{kb_id}")
        if not 200 <= cleanup.status_code < 300:
            raise RuntimeError(
                f"cleanup delete non-2xx: {cleanup.status_code} {cleanup.text[:500]}"
            )
    remaining = await _list_owned_records(client, owned_name)
    if remaining:
        raise RuntimeError(f"cleanup verification found {len(remaining)} owned records")
    return sorted(owned_ids)


async def _run_format(
    client: httpx.AsyncClient,
    kind: str,
    path: Path,
    timeout: float,
    probes: dict[str, list[dict[str, Any]]],
) -> dict[str, Any]:
    started_at = _utc_now()
    started = time.perf_counter()
    result: dict[str, Any] = {
        "started_at": started_at,
        "status": "failed",
        "document_count": 0,
        "cleanup_status": "not_started",
        "error": None,
    }
    kb_id = None
    owned_name = f"document-concurrency-{kind}-{uuid.uuid4().hex[:12]}"
    result["owned_name"] = owned_name
    try:
        with path.open("rb") as stream:
            response = await client.post(
                "/api/v1/knowledge/upload",
                data={"name": owned_name, "category": "manual"},
                files={"file": (path.name, stream, "application/octet-stream")},
            )
        result["upload_status_code"] = response.status_code
        if not 200 <= response.status_code < 300:
            raise RuntimeError(f"upload non-2xx: {response.status_code} {response.text[:500]}")
        kb_id = (response.json().get("data") or {}).get("id")
        if not kb_id:
            raise RuntimeError("upload response did not include a task ID")
        result["knowledge_base_id"] = kb_id

        health_tasks = [
            asyncio.create_task(_probe(client, "/health", kind, "health"))
            for _ in range(HEALTH_SAMPLES_PER_FORMAT)
        ]
        project_tasks = [
            asyncio.create_task(
                _probe(client, "/api/v1/novel-writer/projects", kind, "projects")
            )
            for _ in range(PROJECT_SAMPLES_PER_FORMAT)
        ]
        poll_task = asyncio.create_task(_poll_completion(client, kb_id, timeout))
        try:
            detail = await poll_task
        finally:
            # Probe evidence must survive parse/poll failures.  Never let one
            # terminal task discard already completed raw latency samples.
            health_samples = await asyncio.gather(*health_tasks)
            project_samples = await asyncio.gather(*project_tasks)
            probes["health"].extend(health_samples)
            probes["projects"].extend(project_samples)
        result["document_count"] = int(detail.get("document_count") or 0)
        result["status"] = "completed"
    except Exception as exc:
        result["error"] = f"{type(exc).__name__}: {exc}"
    finally:
        try:
            cleaned_ids = await _cleanup_owned_records(client, owned_name, kb_id)
            result["cleaned_knowledge_base_ids"] = cleaned_ids
            result["cleanup_status"] = "completed"
        except Exception as exc:
            result["cleanup_status"] = "failed"
            result["cleanup_error"] = f"{type(exc).__name__}: {exc}"
        result["duration_seconds"] = round(time.perf_counter() - started, 6)
    return result


async def run_verification(args: argparse.Namespace) -> dict[str, Any]:
    started_at = _utc_now()
    started = time.perf_counter()
    probes: dict[str, list[dict[str, Any]]] = {"health": [], "projects": []}
    report: dict[str, Any] = {
        "schema_version": 1,
        "started_at": started_at,
        "base_url": args.base_url,
        "sample_requirements": {
            "health_per_format": HEALTH_SAMPLES_PER_FORMAT,
            "projects_per_format": PROJECT_SAMPLES_PER_FORMAT,
        },
        "formats": {},
        "probes": probes,
    }
    headers = {"Authorization": f"Bearer {args.token}"}
    with tempfile.TemporaryDirectory(prefix="document-concurrency-") as temporary:
        samples = generate_samples(Path(temporary), scale=args.scale)
        report["sample_validation"] = validate_generated_samples(samples)
        async with httpx.AsyncClient(
            base_url=args.base_url.rstrip("/"), headers=headers, timeout=args.request_timeout
        ) as client:
            for kind in FORMATS:
                report["formats"][kind] = await _run_format(
                    client, kind, samples[kind], args.processing_timeout, probes
                )
    report["finished_at"] = _utc_now()
    report["total_duration_seconds"] = round(time.perf_counter() - started, 6)
    report["errors"] = validate_report(report)
    return report


def _parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--base-url", default=os.getenv("DOCUMENT_VERIFY_BASE_URL"))
    parser.add_argument("--token", default=os.getenv("DOCUMENT_VERIFY_TOKEN"))
    parser.add_argument("--json-output", type=Path, default=DEFAULT_JSON_OUTPUT)
    parser.add_argument("--markdown-output", type=Path, default=DEFAULT_MARKDOWN_OUTPUT)
    parser.add_argument("--scale", type=int, default=3)
    parser.add_argument("--processing-timeout", type=float, default=360.0)
    parser.add_argument("--request-timeout", type=float, default=30.0)
    return parser


def main(argv: list[str] | None = None) -> int:
    args = _parser().parse_args(argv)
    if not args.base_url or not args.token:
        print(
            "ERROR: --base-url/DOCUMENT_VERIFY_BASE_URL and "
            "--token/DOCUMENT_VERIFY_TOKEN are required; verification is fail-closed.",
            file=sys.stderr,
        )
        return 2
    try:
        report = asyncio.run(run_verification(args))
        args.json_output.parent.mkdir(parents=True, exist_ok=True)
        args.json_output.write_text(
            json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8"
        )
        persisted_report = json.loads(args.json_output.read_text(encoding="utf-8"))
        summary = build_summary(persisted_report)
        args.markdown_output.parent.mkdir(parents=True, exist_ok=True)
        args.markdown_output.write_text(
            render_markdown(persisted_report, summary), encoding="utf-8"
        )
    except Exception as exc:
        print(f"ERROR: {type(exc).__name__}: {exc}", file=sys.stderr)
        return 1
    print(f"Raw JSON: {args.json_output}")
    print(f"Markdown: {args.markdown_output}")
    return 1 if report.get("errors") else 0


if __name__ == "__main__":
    raise SystemExit(main())
