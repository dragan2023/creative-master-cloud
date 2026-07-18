"""Knowledge upload integration through handler, worker, parser, and DB state."""

from __future__ import annotations

import asyncio
from io import BytesIO
import importlib
from pathlib import Path
from types import SimpleNamespace

import pytest
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker

from app.models import (
    KnowledgeBase,
    KnowledgeBaseStatus,
    SystemConfig,
    User,
    UserAPIKey,
)
from app.tools.file_parser import FileParser


upload_module = importlib.import_module("app.api.v1.endpoints.knowledge._upload")
state_module = importlib.import_module("app.api.v1.endpoints.knowledge._state")


def _pdf_bytes(path: Path) -> bytes:
    from pypdf import PdfWriter
    from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

    writer = PdfWriter()
    for text in ("Knowledge PDF page one", "Knowledge PDF page two"):
        page = writer.add_blank_page(width=612, height=792)
        font = DictionaryObject(
            {
                NameObject("/Type"): NameObject("/Font"),
                NameObject("/Subtype"): NameObject("/Type1"),
                NameObject("/BaseFont"): NameObject("/Helvetica"),
            }
        )
        page[NameObject("/Resources")] = DictionaryObject(
            {
                NameObject("/Font"): DictionaryObject(
                    {NameObject("/F1"): writer._add_object(font)}
                )
            }
        )
        content = DecodedStreamObject()
        content.set_data(f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode("ascii"))
        page[NameObject("/Contents")] = writer._add_object(content)
    with path.open("wb") as stream:
        writer.write(stream)
    return path.read_bytes()


def _docx_bytes(path: Path) -> bytes:
    from docx import Document

    document = Document()
    document.add_heading("知识文档", level=1)
    document.add_paragraph("真实处理链路。")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "字段"
    table.cell(0, 1).text = "值"
    document.save(path)
    return path.read_bytes()


def _xlsx_bytes(path: Path) -> bytes:
    from openpyxl import Workbook

    workbook = Workbook()
    workbook.active.append(["知识", 42, None])
    workbook.create_sheet("第二页")["A1"] = "补充内容"
    workbook.save(path)
    return path.read_bytes()


def _valid_samples(tmp_path: Path) -> dict[str, bytes]:
    return {
        ".pdf": _pdf_bytes(tmp_path / "knowledge.pdf"),
        ".docx": _docx_bytes(tmp_path / "knowledge.docx"),
        ".xlsx": _xlsx_bytes(tmp_path / "knowledge.xlsx"),
        ".txt": "知识上传中文内容。\n第二行。".encode("utf-8"),
    }


class PersistingUploadDB:
    """Request-session adapter that persists ORM records to the test DB."""

    def __init__(self, session_factory):
        self.session_factory = session_factory
        self.pending = None

    def add(self, value):
        self.pending = value

    async def commit(self):
        if self.pending is None:
            return
        with self.session_factory() as session:
            session.add(self.pending)
            session.commit()
            session.refresh(self.pending)
            session.expunge(self.pending)

    async def refresh(self, _value):
        return None

    async def rollback(self):
        return None


class InMemoryRetrieval:
    def __init__(self):
        self.collections: dict[str, list[str]] = {}

    async def add_documents_batch(self, collection_name, documents, metadatas):
        assert len(documents) == len(metadatas)
        assert all(metadata["knowledge_base_id"] > 0 for metadata in metadatas)
        self.collections[collection_name] = list(documents)


async def _wait_for_terminal(session_factory, kb_id: int):
    for _ in range(200):
        with session_factory() as session:
            record = session.get(KnowledgeBase, kb_id)
            if record and record.status in {KnowledgeBaseStatus.READY, KnowledgeBaseStatus.FAILED}:
                session.expunge(record)
                return record
        await asyncio.sleep(0.025)
    raise AssertionError(f"knowledge upload {kb_id} did not reach a terminal state")


@pytest.mark.asyncio
async def test_real_handler_completes_four_formats_and_fails_corrupt_xlsx(
    monkeypatch,
    tmp_path: Path,
):
    database_path = tmp_path / "knowledge-upload.db"
    database_url = f"sqlite:///{database_path.as_posix()}"
    engine = create_engine(database_url, connect_args={"check_same_thread": False})
    for table in (User.__table__, UserAPIKey.__table__, SystemConfig.__table__, KnowledgeBase.__table__):
        table.create(engine, checkfirst=True)
    Session = sessionmaker(bind=engine, expire_on_commit=False)

    progress_store = {}
    retrieval = InMemoryRetrieval()
    monkeypatch.setattr(upload_module, "SYNC_DATABASE_URL", database_url)
    monkeypatch.setattr(upload_module, "get_file_parser", lambda: FileParser(use_preprocessor=False))
    monkeypatch.setattr(upload_module, "get_knowledge_retrieval_tool", lambda: retrieval)
    monkeypatch.setattr(upload_module, "kb_processing_progress", progress_store)
    monkeypatch.setattr(state_module, "kb_processing_progress", progress_store)
    monkeypatch.setattr(state_module, "_sync_update_kb_progress", lambda *_args: None)
    monkeypatch.setattr(
        upload_module,
        "settings",
        SimpleNamespace(
            ALLOWED_EXTENSIONS={".pdf", ".docx", ".xlsx", ".txt"},
            MAX_UPLOAD_SIZE=5 * 1024 * 1024,
            get_chroma_dir=lambda: str(tmp_path / "chroma"),
        ),
    )

    completed = []
    try:
        for index, (extension, content) in enumerate(_valid_samples(tmp_path).items()):
            response = await upload_module.upload_knowledge_base_handler(
                background_tasks=None,
                name=f"real-{extension[1:]}",
                file=SimpleNamespace(
                    filename=f"sample{extension}", size=len(content), file=BytesIO(content)
                ),
                description="integration",
                category="manual",
                current_user=SimpleNamespace(id=900 + index),
                db=PersistingUploadDB(Session),
            )
            assert response.status == KnowledgeBaseStatus.PROCESSING
            record = await _wait_for_terminal(Session, response.id)
            assert record.status == KnowledgeBaseStatus.READY
            assert record.document_count > 0
            assert retrieval.collections[record.collection_name]
            assert progress_store[record.id]["status"] == "completed"
            completed.append(record)

        corrupt = b"not an xlsx package"
        response = await upload_module.upload_knowledge_base_handler(
            background_tasks=None,
            name="corrupt-xlsx",
            file=SimpleNamespace(
                filename="corrupt.xlsx", size=len(corrupt), file=BytesIO(corrupt)
            ),
            description="integration",
            category="manual",
            current_user=SimpleNamespace(id=999),
            db=PersistingUploadDB(Session),
        )
        failed = await _wait_for_terminal(Session, response.id)
        assert failed.status == KnowledgeBaseStatus.FAILED
        assert failed.document_count == 0
        assert failed.collection_name not in retrieval.collections
        assert progress_store[failed.id]["status"] == "failed"
        assert progress_store[failed.id]["error"].startswith("KB-PARSE-001:")
    finally:
        for task in list(state_module.kb_processing_tasks.values()):
            future = task.get("future")
            if future:
                future.result(timeout=5)
        engine.dispose()

    assert len(completed) == 4
