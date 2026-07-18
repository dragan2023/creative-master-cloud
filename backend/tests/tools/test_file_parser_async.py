"""Real-document compatibility and event-loop safety tests for FileParser."""

from __future__ import annotations

import asyncio
import importlib
from pathlib import Path
import threading

import pytest

from app.tools.file_parser import FileParser


def _write_pdf(path: Path) -> None:
    from pypdf import PdfReader, PdfWriter
    from pypdf.generic import (
        DecodedStreamObject,
        DictionaryObject,
        NameObject,
    )

    writer = PdfWriter()
    for text in ("First PDF page", "Second PDF page"):
        page = writer.add_blank_page(width=612, height=792)
        font = DictionaryObject(
            {
                NameObject("/Type"): NameObject("/Font"),
                NameObject("/Subtype"): NameObject("/Type1"),
                NameObject("/BaseFont"): NameObject("/Helvetica"),
            }
        )
        resources = DictionaryObject(
            {
                NameObject("/Font"): DictionaryObject(
                    {NameObject("/F1"): writer._add_object(font)}
                )
            }
        )
        stream = DecodedStreamObject()
        stream.set_data(
            f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode("ascii")
        )
        page[NameObject("/Resources")] = resources
        page[NameObject("/Contents")] = writer._add_object(stream)
    with path.open("wb") as output:
        writer.write(output)

    reopened = PdfReader(path)
    assert len(reopened.pages) == 2
    assert [page.extract_text().strip() for page in reopened.pages] == [
        "First PDF page",
        "Second PDF page",
    ]


def _write_docx(path: Path) -> None:
    from docx import Document

    document = Document()
    document.add_heading("兼容性标题", level=1)
    document.add_paragraph("正文段落。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "名称"
    table.cell(0, 1).text = "数量"
    table.cell(1, 0).text = "苹果"
    table.cell(1, 1).text = "3"
    document.save(path)

    reopened = Document(path)
    assert len(reopened.tables) == 1
    assert reopened.paragraphs[0].text == "兼容性标题"


def _write_xlsx(path: Path) -> None:
    from openpyxl import Workbook, load_workbook

    workbook = Workbook()
    first = workbook.active
    first.title = "主表"
    first.append(["项目", "数量", None])
    first.append(["铅笔", 12, None])
    second = workbook.create_sheet("备注")
    second["A1"] = "第二工作表"
    second["C2"] = 7.5
    workbook.save(path)

    reopened = load_workbook(path, read_only=True, data_only=True)
    try:
        assert reopened.sheetnames == ["主表", "备注"]
        assert reopened["主表"]["B2"].value == 12
    finally:
        reopened.close()


def _write_txt(path: Path) -> None:
    path.write_text("中文测试。\n第二行！", encoding="utf-8")
    assert path.read_text(encoding="utf-8") == "中文测试。\n第二行！"


@pytest.fixture
def real_documents(tmp_path: Path) -> dict[str, Path]:
    paths = {
        "pdf": tmp_path / "sample.pdf",
        "docx": tmp_path / "sample.docx",
        "xlsx": tmp_path / "sample.xlsx",
        "txt": tmp_path / "sample.txt",
    }
    _write_pdf(paths["pdf"])
    _write_docx(paths["docx"])
    _write_xlsx(paths["xlsx"])
    _write_txt(paths["txt"])
    return paths


@pytest.mark.asyncio
@pytest.mark.parametrize(
    ("kind", "metadata"),
    [
        ("pdf", {"file_type": "pdf", "page_count": 2}),
        (
            "docx",
            {"file_type": "docx", "paragraph_count": 2, "table_count": 1},
        ),
        ("xlsx", {"file_type": "xlsx", "sheet_count": 2}),
        ("txt", {"file_type": "txt", "encoding": "utf-8"}),
    ],
)
async def test_sync_and_public_async_results_match_field_for_field(
    real_documents: dict[str, Path], kind: str, metadata: dict[str, object]
) -> None:
    from app.tools.document_extractors import extract_document

    path = real_documents[kind]
    sync_result = extract_document(str(path))
    async_result = await FileParser(use_preprocessor=False).parse(str(path))

    assert async_result == sync_result
    assert sync_result["content"].strip()
    assert sync_result["metadata"] == {
        "file_path": str(path),
        **metadata,
        "char_count": len(sync_result["content"]),
    }


@pytest.mark.asyncio
async def test_real_document_content_is_extracted_completely(
    real_documents: dict[str, Path],
) -> None:
    parser = FileParser(use_preprocessor=False)
    results = {kind: await parser.parse(str(path)) for kind, path in real_documents.items()}

    assert results["pdf"]["content"] == "First PDF page\n\nSecond PDF page"
    assert results["docx"]["content"] == (
        "兼容性标题\n正文段落。\n\n表格内容:\n名称 | 数量\n苹果 | 3"
    )
    assert results["xlsx"]["content"] == (
        "【工作表: 主表】\n项目 | 数量\n铅笔 | 12\n\n"
        "【工作表: 备注】\n第二工作表\n7.5"
    )
    assert results["txt"]["content"] == "中文测试。\n第二行！"


@pytest.mark.asyncio
@pytest.mark.parametrize("extension", [".pdf", ".docx", ".xlsx", ".txt"])
async def test_corrupt_documents_fail_explicitly(tmp_path: Path, extension: str) -> None:
    from app.tools.document_extractors import extract_document

    path = tmp_path / f"corrupt{extension}"
    path.write_bytes(b"" if extension == ".txt" else b"not a structurally valid document")

    sync_result = extract_document(str(path))
    async_result = await FileParser(use_preprocessor=False).parse(str(path))

    assert async_result == sync_result
    assert set(sync_result) == {"error"}
    assert "DOC-EXTRACT-" in sync_result["error"]


@pytest.mark.asyncio
async def test_empty_but_valid_document_is_rejected(tmp_path: Path) -> None:
    from docx import Document
    from app.tools.document_extractors import extract_document

    path = tmp_path / "empty.docx"
    Document().save(path)
    result = extract_document(str(path))

    assert set(result) == {"error"}
    assert "DOC-EXTRACT-EMPTY" in result["error"]


def test_legacy_xls_is_not_advertised_or_parsed() -> None:
    parser = FileParser(use_preprocessor=False)
    assert parser.is_supported("legacy.xls") is False


def test_application_and_preprocessor_extension_contracts_match() -> None:
    from app.core.config import get_settings
    from app.tools.doc_preprocessor import DocumentPreprocessor

    settings = get_settings()
    preprocessor = DocumentPreprocessor(settings=settings)
    for extension in (".pdf", ".docx", ".xlsx", ".txt"):
        assert extension in settings.ALLOWED_EXTENSIONS
        assert preprocessor.is_supported(f"document{extension}")
    assert ".xls" not in settings.ALLOWED_EXTENSIONS
    assert preprocessor.is_supported("document.xls") is False


@pytest.mark.asyncio
async def test_public_async_parse_keeps_event_loop_responsive(monkeypatch, tmp_path: Path) -> None:
    started = threading.Event()
    release = threading.Event()
    heartbeat_ran = asyncio.Event()

    def blocking_extract(path: str):
        started.set()
        release.wait(2)
        return {
            "content": "released",
            "metadata": {
                "file_path": path,
                "file_type": "txt",
                "char_count": 8,
                "encoding": "utf-8",
            },
        }

    file_parser_module = importlib.import_module("app.tools.file_parser")
    monkeypatch.setattr(file_parser_module, "extract_document", blocking_extract)
    path = tmp_path / "heartbeat.txt"
    path.write_text("input", encoding="utf-8")

    task = asyncio.create_task(FileParser(use_preprocessor=False).parse(str(path)))
    assert await asyncio.to_thread(started.wait, 1)
    await asyncio.sleep(0)
    heartbeat_ran.set()
    assert heartbeat_ran.is_set()
    release.set()
    assert (await task)["content"] == "released"
