"""
导出服务
支持多种格式的项目导出
"""
import os
import json
from typing import Dict, Any, Optional, List
from datetime import datetime

from app.core.logger import get_logger
from app.models import NovelProject, NovelChapter


class NovelExporter:
    """小说/剧本导出服务

    支持格式：
    - txt: 纯文本
    - md: Markdown格式
    - docx: Word文档
    - epub: 电子书格式
    """

    def __init__(self):
        self.logger = get_logger("novel_exporter")

    async def export_project(
        self,
        project: NovelProject,
        chapters: List[NovelChapter],
        format: str = "txt",
        include_metadata: bool = False,
        chapter_range: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        导出项目

        Args:
            project: 项目对象
            chapters: 章节列表
            format: 导出格式
            include_metadata: 是否包含元数据
            chapter_range: 章节范围（如 "1-10"）

        Returns:
            导出结果
        """
        result = {
            "success": False,
            "file_path": None,
            "file_name": None,
            "file_size": 0,
            "format": format,
            "error_message": None
        }

        try:
            # 解析章节范围
            target_chapters = self._parse_chapter_range(
                chapters, chapter_range)

            if not target_chapters:
                result["error_message"] = "没有可导出的章节"
                return result

            # 根据格式选择导出方法
            export_methods = {
                "txt": self._export_txt,
                "md": self._export_markdown,
                "docx": self._export_docx,
                "epub": self._export_epub
            }

            export_method = export_methods.get(format, self._export_txt)

            # 执行导出
            export_result = await export_method(
                project, target_chapters, include_metadata
            )

            if export_result.get("success"):
                result["success"] = True
                result["file_path"] = export_result["file_path"]
                result["file_name"] = export_result["file_name"]
                result["file_size"] = export_result.get("file_size", 0)
            else:
                result["error_message"] = export_result.get(
                    "error_message", "导出失败")

            return result

        except Exception as e:
            self.logger.error(f"导出项目失败: {str(e)}")
            result["error_message"] = str(e)
            return result

    def _parse_chapter_range(
        self,
        chapters: List[NovelChapter],
        chapter_range: Optional[str]
    ) -> List[NovelChapter]:
        """解析章节范围"""
        if not chapter_range:
            return [c for c in chapters if c.final_content]

        try:
            if "-" in chapter_range:
                start, end = map(int, chapter_range.split("-"))
                return [
                    c for c in chapters
                    if start <= c.chapter_number <= end and c.final_content
                ]
            else:
                # 单个章节
                chapter_num = int(chapter_range)
                return [c for c in chapters if c.chapter_number == chapter_num and c.final_content]
        except:
            return [c for c in chapters if c.final_content]

    async def _export_txt(
        self,
        project: NovelProject,
        chapters: List[NovelChapter],
        include_metadata: bool
    ) -> Dict[str, Any]:
        """导出为TXT格式"""
        try:
            # 创建输出目录
            export_dir = os.path.join("data", "exports")
            os.makedirs(export_dir, exist_ok=True)

            # 生成文件名
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            file_name = f"{project.title}_{timestamp}.txt"
            file_path = os.path.join(export_dir, file_name)

            # 构建内容
            lines = []

            # 标题
            lines.append(f"《{project.title}》")
            lines.append("")
            lines.append(
                f"类型: {'小说' if project.project_type.value == 'novel' else '剧本'}")
            lines.append(f"总字数: {sum(c.word_count for c in chapters)}")
            lines.append("")
            lines.append("=" * 50)
            lines.append("")

            # 元数据
            if include_metadata:
                lines.append("【项目信息】")
                lines.append(f"题材: {project.genre or '未指定'}")
                lines.append(f"目标平台: {project.target_platform or '未指定'}")
                lines.append(f"创建时间: {project.created_at}")
                lines.append("")
                lines.append("=" * 50)
                lines.append("")

            # 章节
            for chapter in sorted(chapters, key=lambda c: c.chapter_number):
                lines.append(
                    f"第{chapter.chapter_number}章 {chapter.chapter_title or ''}")
                lines.append("")

                if include_metadata and chapter.chapter_metadata:
                    lines.append("【本章信息】")
                    metadata = chapter.chapter_metadata
                    if metadata.get("chapter_role"):
                        lines.append(f"章节定位: {metadata['chapter_role']}")
                    if metadata.get("chapter_purpose"):
                        lines.append(f"核心作用: {metadata['chapter_purpose']}")
                    lines.append("")

                lines.append(chapter.final_content or "")
                lines.append("")
                lines.append("-" * 30)
                lines.append("")

            # 写入文件
            content = "\n".join(lines)
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)

            return {
                "success": True,
                "file_path": file_path,
                "file_name": file_name,
                "file_size": len(content.encode('utf-8'))
            }

        except Exception as e:
            self.logger.error(f"导出TXT失败: {str(e)}")
            return {"success": False, "error_message": str(e)}

    async def _export_markdown(
        self,
        project: NovelProject,
        chapters: List[NovelChapter],
        include_metadata: bool
    ) -> Dict[str, Any]:
        """导出为Markdown格式"""
        try:
            export_dir = os.path.join("data", "exports")
            os.makedirs(export_dir, exist_ok=True)

            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            file_name = f"{project.title}_{timestamp}.md"
            file_path = os.path.join(export_dir, file_name)

            lines = []

            # 标题
            lines.append(f"# 《{project.title}》")
            lines.append("")
            lines.append(
                f"> 类型: {'小说' if project.project_type.value == 'novel' else '剧本'} | 总字数: {sum(c.word_count for c in chapters)}")
            lines.append("")

            # 元数据
            if include_metadata:
                lines.append("## 📋 项目信息")
                lines.append("")
                lines.append(f"- **题材**: {project.genre or '未指定'}")
                lines.append(f"- **目标平台**: {project.target_platform or '未指定'}")
                lines.append(f"- **创建时间**: {project.created_at}")
                lines.append("")

            # 目录
            lines.append("## 📚 目录")
            lines.append("")
            for chapter in sorted(chapters, key=lambda c: c.chapter_number):
                anchor = f"chapter-{chapter.chapter_number}"
                lines.append(
                    f"- [第{chapter.chapter_number}章 {chapter.chapter_title or ''}](#{anchor})")
            lines.append("")

            # 章节
            lines.append("## 📖 正文")
            lines.append("")

            for chapter in sorted(chapters, key=lambda c: c.chapter_number):
                lines.append(
                    f"### 第{chapter.chapter_number}章 {chapter.chapter_title or ''} {{#chapter-{chapter.chapter_number}}}")
                lines.append("")

                if include_metadata and chapter.chapter_metadata:
                    metadata = chapter.chapter_metadata
                    lines.append("> **章节信息**")
                    if metadata.get("chapter_role"):
                        lines.append(f"> - 定位: {metadata['chapter_role']}")
                    if metadata.get("chapter_purpose"):
                        lines.append(f"> - 作用: {metadata['chapter_purpose']}")
                    lines.append("")

                # 处理段落
                paragraphs = (chapter.final_content or "").split("\n")
                for para in paragraphs:
                    if para.strip():
                        lines.append(para.strip())
                        lines.append("")

            content = "\n".join(lines)
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)

            return {
                "success": True,
                "file_path": file_path,
                "file_name": file_name,
                "file_size": len(content.encode('utf-8'))
            }

        except Exception as e:
            self.logger.error(f"导出Markdown失败: {str(e)}")
            return {"success": False, "error_message": str(e)}

    async def _export_docx(
        self,
        project: NovelProject,
        chapters: List[NovelChapter],
        include_metadata: bool
    ) -> Dict[str, Any]:
        """导出为Word文档"""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            export_dir = os.path.join("data", "exports")
            os.makedirs(export_dir, exist_ok=True)

            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            file_name = f"{project.title}_{timestamp}.docx"
            file_path = os.path.join(export_dir, file_name)

            doc = Document()

            # 标题
            title = doc.add_heading(f"《{project.title}》", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 项目信息
            doc.add_paragraph(
                f"类型: {'小说' if project.project_type.value == 'novel' else '剧本'}")
            doc.add_paragraph(f"总字数: {sum(c.word_count for c in chapters)}")

            if include_metadata:
                doc.add_paragraph(f"题材: {project.genre or '未指定'}")
                doc.add_paragraph(f"目标平台: {project.target_platform or '未指定'}")

            doc.add_paragraph("")  # 空行

            # 章节
            for chapter in sorted(chapters, key=lambda c: c.chapter_number):
                # 章节标题
                doc.add_heading(
                    f"第{chapter.chapter_number}章 {chapter.chapter_title or ''}", level=1)

                if include_metadata and chapter.chapter_metadata:
                    metadata = chapter.chapter_metadata
                    info_para = doc.add_paragraph()
                    if metadata.get("chapter_role"):
                        info_para.add_run(
                            f"[章节定位: {metadata['chapter_role']}] ")
                    if metadata.get("chapter_purpose"):
                        info_para.add_run(
                            f"[核心作用: {metadata['chapter_purpose']}]")

                # 章节内容
                paragraphs = (chapter.final_content or "").split("\n")
                for para in paragraphs:
                    if para.strip():
                        p = doc.add_paragraph(para.strip())
                        p.paragraph_format.first_line_indent = Inches(0.3)

                doc.add_paragraph("")  # 空行

            doc.save(file_path)

            file_size = os.path.getsize(file_path)

            return {
                "success": True,
                "file_path": file_path,
                "file_name": file_name,
                "file_size": file_size
            }

        except ImportError:
            return {
                "success": False,
                "error_message": "需要安装 python-docx 库: pip install python-docx"
            }
        except Exception as e:
            self.logger.error(f"导出DOCX失败: {str(e)}")
            return {"success": False, "error_message": str(e)}

    async def _export_epub(
        self,
        project: NovelProject,
        chapters: List[NovelChapter],
        include_metadata: bool
    ) -> Dict[str, Any]:
        """导出为EPUB格式"""
        try:
            from ebooklib import epub

            export_dir = os.path.join("data", "exports")
            os.makedirs(export_dir, exist_ok=True)

            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            file_name = f"{project.title}_{timestamp}.epub"
            file_path = os.path.join(export_dir, file_name)

            book = epub.EpubBook()

            # 设置元数据
            book.set_identifier(f"novel_{project.id}_{timestamp}")
            book.set_title(project.title)
            book.set_language('zh')
            book.add_author('AI创作')

            # 创建章节
            epub_chapters = []
            toc = []

            for chapter in sorted(chapters, key=lambda c: c.chapter_number):
                chapter_file = f"chapter_{chapter.chapter_number}.xhtml"
                epub_chapter = epub.EpubHtml(
                    title=f"第{chapter.chapter_number}章 {chapter.chapter_title or ''}",
                    file_name=chapter_file,
                    lang='zh'
                )

                # 构建HTML内容
                content_html = f"<h1>第{chapter.chapter_number}章 {chapter.chapter_title or ''}</h1>"

                if include_metadata and chapter.chapter_metadata:
                    content_html += "<div class='metadata'>"
                    if chapter.chapter_metadata.get("chapter_role"):
                        content_html += f"<p>章节定位: {chapter.chapter_metadata['chapter_role']}</p>"
                    content_html += "</div>"

                # 添加正文
                paragraphs = (chapter.final_content or "").split("\n")
                for para in paragraphs:
                    if para.strip():
                        content_html += f"<p>{para.strip()}</p>"

                epub_chapter.content = content_html
                book.add_item(epub_chapter)
                epub_chapters.append(epub_chapter)
                toc.append(epub_chapter)

            # 设置目录
            book.toc = tuple(toc)

            # 添加导航文件
            book.add_item(epub.EpubNcx())
            book.add_item(epub.EpubNav())

            # 设置spine
            book.spine = ['nav'] + epub_chapters

            # 写入文件
            epub.write_epub(file_path, book, {})

            file_size = os.path.getsize(file_path)

            return {
                "success": True,
                "file_path": file_path,
                "file_name": file_name,
                "file_size": file_size
            }

        except ImportError:
            return {
                "success": False,
                "error_message": "需要安装 ebooklib 库: pip install ebooklib"
            }
        except Exception as e:
            self.logger.error(f"导出EPUB失败: {str(e)}")
            return {"success": False, "error_message": str(e)}
