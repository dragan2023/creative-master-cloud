"""Synchronous document text extraction without async or database dependencies."""

from __future__ import annotations

from pathlib import Path
from typing import Any, Callable


SUPPORTED_DOCUMENT_EXTENSIONS = frozenset(
    {".pdf", ".docx", ".doc", ".xlsx", ".txt", ".md"}
)
TEXT_ENCODINGS = ("utf-8", "gbk", "gb2312", "gb18030", "utf-16", "latin-1")


def _pdf(path: str) -> dict[str, Any]:
    from pypdf import PdfReader

    reader = PdfReader(path)
    page_text = []
    for page in reader.pages:
        extracted = page.extract_text()
        if extracted and extracted.strip():
            page_text.append(extracted.strip())
    content = "\n\n".join(page_text)
    return {
        "content": content,
        "metadata": {
            "file_path": path,
            "file_type": "pdf",
            "page_count": len(reader.pages),
            "char_count": len(content),
        },
    }


def _docx(path: str) -> dict[str, Any]:
    from docx import Document

    document = Document(path)
    paragraphs = [item.text for item in document.paragraphs if item.text.strip()]
    table_rows = []
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                table_rows.append(" | ".join(values))
    content = "\n".join(paragraphs)
    if table_rows:
        content += ("\n\n" if content else "") + "表格内容:\n" + "\n".join(table_rows)
    return {
        "content": content,
        "metadata": {
            "file_path": path,
            "file_type": "docx",
            "paragraph_count": len(paragraphs),
            "table_count": len(document.tables),
            "char_count": len(content),
        },
    }


def _xlsx(path: str) -> dict[str, Any]:
    from openpyxl import load_workbook

    workbook = load_workbook(path, read_only=True, data_only=True)
    try:
        sheet_names = list(workbook.sheetnames)
        sheets = []
        for sheet_name in sheet_names:
            rows = []
            for row in workbook[sheet_name].iter_rows(values_only=True):
                values = [
                    str(value)
                    for value in row
                    if value is not None and str(value).strip()
                ]
                if values:
                    rows.append(" | ".join(values))
            if rows:
                sheets.append(f"【工作表: {sheet_name}】\n" + "\n".join(rows))
        content = "\n\n".join(sheets)
    finally:
        workbook.close()
    return {
        "content": content,
        "metadata": {
            "file_path": path,
            "file_type": "xlsx",
            "sheet_count": len(sheet_names),
            "char_count": len(content),
        },
    }


def _text(path: str) -> dict[str, Any]:
    last_decode_error: UnicodeDecodeError | None = None
    for encoding in TEXT_ENCODINGS:
        try:
            with open(path, "r", encoding=encoding) as stream:
                content = stream.read()
            return {
                "content": content,
                "metadata": {
                    "file_path": path,
                    "file_type": "txt",
                    "char_count": len(content),
                    "encoding": encoding,
                },
            }
        except UnicodeDecodeError as error:
            last_decode_error = error
    if last_decode_error is not None:
        raise last_decode_error
    raise ValueError("no text encoding was attempted")


_EXTRACTORS: dict[str, Callable[[str], dict[str, Any]]] = {
    ".pdf": _pdf,
    ".docx": _docx,
    ".doc": _docx,
    ".xlsx": _xlsx,
    ".txt": _text,
    ".md": _text,
}


def extract_document(file_path: str) -> dict[str, Any]:
    """Extract a document while preserving FileParser's public result shape."""

    path = str(file_path)
    suffix = Path(path).suffix.lower()
    if not Path(path).exists():
        return {"error": f"DOC-EXTRACT-NOT-FOUND: 文件不存在: {path}"}
    extractor = _EXTRACTORS.get(suffix)
    if extractor is None:
        return {"error": f"DOC-EXTRACT-UNSUPPORTED: 不支持的文件格式: {suffix}"}
    try:
        result = extractor(path)
    except ImportError as error:
        return {"error": f"DOC-EXTRACT-DEPENDENCY: {suffix} 解析依赖缺失: {error}"}
    except Exception as error:
        return {"error": f"DOC-EXTRACT-INVALID: {suffix} 文件损坏或格式无效: {error}"}
    if not result.get("content", "").strip():
        return {"error": f"DOC-EXTRACT-EMPTY: {suffix} 文件未提取到有效文本"}
    return result


__all__ = ["SUPPORTED_DOCUMENT_EXTENSIONS", "TEXT_ENCODINGS", "extract_document"]
