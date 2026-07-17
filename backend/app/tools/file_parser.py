"""
文件解析工具
支持 PDF、Word、TXT 等格式的文件解析
集成文档预处理流水线 (Cleaner -> Filter -> Refiner)

@date: 2026-04-02
@version: v3.0.0
@author: 周金磊
@contact: QQ：7527149（添加时请说明来意）
"""
from typing import List, Dict, Any, Optional, Callable
from pathlib import Path
import os
import re

from app.core.config import get_settings


class FileParser:
    """文件解析器"""

    def __init__(self, use_preprocessor: bool = None):
        """
        初始化文件解析器

        Args:
            use_preprocessor: 是否使用预处理器（None时使用配置文件设置）
        """
        self.supported_extensions = {".pdf", ".docx", ".doc", ".txt", ".md", ".xlsx", ".xls"}
        self.settings = get_settings()

        # 决定是否使用预处理器
        if use_preprocessor is None:
            use_preprocessor = self.settings.DOC_PREPROCESSOR_ENABLED

        self.use_preprocessor = use_preprocessor
        self._preprocessor = None

    def _get_preprocessor(self, config=None):
        """延迟加载预处理器"""
        if self._preprocessor is None and self.use_preprocessor:
            from app.tools.doc_preprocessor import DocumentPreprocessor
            self._preprocessor = DocumentPreprocessor(
                self.settings, config=config)
        elif config and self._preprocessor:
            # 更新配置
            self._preprocessor.config = config
        return self._preprocessor

    def is_supported(self, file_path: str) -> bool:
        """
        检查文件是否支持

        Args:
            file_path: 文件路径

        Returns:
            是否支持
        """
        ext = Path(file_path).suffix.lower()
        return ext in self.supported_extensions

    async def parse(self, file_path: str) -> Dict[str, Any]:
        """
        解析文件

        Args:
            file_path: 文件路径

        Returns:
            解析结果，包含文本内容和元数据
        """
        if not os.path.exists(file_path):
            return {"error": f"文件不存在: {file_path}"}

        if not self.is_supported(file_path):
            return {"error": f"不支持的文件格式: {Path(file_path).suffix}"}

        ext = Path(file_path).suffix.lower()

        try:
            if ext == ".pdf":
                return await self._parse_pdf(file_path)
            elif ext in [".docx", ".doc"]:
                return await self._parse_docx(file_path)
            elif ext in [".txt", ".md"]:
                return await self._parse_txt(file_path)
            elif ext in [".xlsx", ".xls"]:
                return await self._parse_xlsx(file_path)
            else:
                return {"error": f"未实现的解析器: {ext}"}
        except Exception as e:
            return {"error": f"解析失败: {str(e)}"}

    async def _parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """
        解析 PDF 文件

        Args:
            file_path: 文件路径

        Returns:
            解析结果
        """
        from pypdf import PdfReader

        reader = PdfReader(file_path)

        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)

        full_text = "\n\n".join(text_parts)

        return {
            "content": full_text,
            "metadata": {
                "file_path": file_path,
                "file_type": "pdf",
                "page_count": len(reader.pages),
                "char_count": len(full_text)
            }
        }

    async def _parse_docx(self, file_path: str) -> Dict[str, Any]:
        """
        解析 Word 文件

        Args:
            file_path: 文件路径

        Returns:
            解析结果
        """
        from docx import Document

        doc = Document(file_path)

        # 提取段落文本
        paragraphs = [
            para.text for para in doc.paragraphs if para.text.strip()]

        # 提取表格文本
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                tables_text.append(row_text)

        full_text = "\n".join(paragraphs)
        if tables_text:
            full_text += "\n\n表格内容:\n" + "\n".join(tables_text)

        return {
            "content": full_text,
            "metadata": {
                "file_path": file_path,
                "file_type": "docx",
                "paragraph_count": len(paragraphs),
                "table_count": len(doc.tables),
                "char_count": len(full_text)
            }
        }

    async def _parse_xlsx(self, file_path: str) -> Dict[str, Any]:
        """
        解析 Excel 文件（.xlsx / .xls）

        Args:
            file_path: 文件路径

        Returns:
            解析结果
        """
        try:
            import openpyxl
        except ImportError:
            return {"error": "缺少 openpyxl 依赖，无法解析 .xlsx 文件。请执行: pip install openpyxl"}

        wb = openpyxl.load_workbook(file_path, data_only=True)
        all_text_parts = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            sheet_parts = [f"【工作表: {sheet_name}】"]

            for row in ws.iter_rows(values_only=True):
                row_values = [str(cell) if cell is not None else "" for cell in row]
                # 跳过完全为空的行
                if any(v.strip() for v in row_values):
                    sheet_parts.append(" | ".join(row_values))

            if len(sheet_parts) > 1:  # 有实际内容
                all_text_parts.append("\n".join(sheet_parts))

        wb.close()
        full_text = "\n\n".join(all_text_parts) if all_text_parts else ""

        return {
            "content": full_text,
            "metadata": {
                "file_path": file_path,
                "file_type": "xlsx",
                "sheet_count": len(wb.sheetnames),
                "char_count": len(full_text)
            }
        }

    async def _parse_txt(self, file_path: str) -> Dict[str, Any]:
        """
        解析 TXT/MD 文件

        Args:
            file_path: 文件路径

        Returns:
            解析结果
        """
        # 尝试多种编码
        encodings = ["utf-8", "gbk", "gb2312", "utf-16", "latin-1"]
        content = None
        used_encoding = None

        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    content = f.read()
                used_encoding = encoding
                break
            except UnicodeDecodeError:
                continue
            except Exception as e:
                return {
                    "error": f"读取文件失败: {str(e)}",
                    "content": ""
                }

        if content is None:
            return {
                "error": "无法识别文件编码，尝试了多种编码均失败",
                "content": ""
            }

        return {
            "content": content,
            "metadata": {
                "file_path": file_path,
                "file_type": "txt",
                "char_count": len(content),
                "encoding": used_encoding
            }
        }

    def split_text(
        self,
        text: str,
        chunk_size: int = 1000,
        overlap: int = 100
    ) -> List[str]:
        """
        将文本分割成块

        Args:
            text: 原始文本
            chunk_size: 块大小（字符数）
            overlap: 重叠大小

        Returns:
            文本块列表
        """
        if len(text) <= chunk_size:
            return [text]

        chunks = []
        start = 0

        while start < len(text):
            end = start + chunk_size

            # 尝试在句子边界分割
            if end < len(text):
                # 查找最近的句子结束符
                last_period = text.rfind("。", start, end)
                last_newline = text.rfind("\n", start, end)

                split_point = max(last_period, last_newline)

                if split_point > start + chunk_size // 2:
                    end = split_point + 1

            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)

            start = end - overlap

        return chunks

    async def parse_and_split(
        self,
        file_path: str,
        chunk_size: int = 1000,
        overlap: int = 100,
        progress_callback: Optional[Callable[[str, int, int], None]] = None,
        llm_provider=None,
        config: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        解析文件并分割成块

        Args:
            file_path: 文件路径
            chunk_size: 块大小（仅在未启用预处理器时使用）
            overlap: 重叠大小（仅在未启用预处理器时使用）
            progress_callback: 进度回调函数 (step_name, progress, step_index)
            llm_provider: LLM提供者（用于摘要压缩）
            config: 用户预处理配置字典

        Returns:
            包含文本块和元数据的结果
        """
        # 优先使用预处理流水线
        preprocessor = self._get_preprocessor(config)
        if preprocessor:
            return await preprocessor.process(file_path, progress_callback, llm_provider)

        # 回退到原有逻辑
        result = await self.parse(file_path)

        if "error" in result:
            return result

        content = result["content"]
        chunks = self.split_text(content, chunk_size, overlap)

        return {
            "chunks": chunks,
            "chunk_count": len(chunks),
            "metadata": result["metadata"]
        }


# 全局文件解析器实例
file_parser = FileParser()


def get_file_parser() -> FileParser:
    """获取文件解析器实例"""
    return file_parser


async def parse_document_file(filename: str, content: bytes) -> str:
    """
    解析文档文件内容（从字节数据）

    Args:
        filename: 文件名
        content: 文件字节内容

    Returns:
        解析后的文本内容
    """
    import tempfile
    import os

    # 获取文件扩展名
    ext = Path(filename).suffix.lower()

    # 创建临时文件
    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp_file:
        tmp_file.write(content)
        tmp_path = tmp_file.name

    try:
        # 使用 FileParser 解析
        parser = FileParser()
        result = await parser.parse(tmp_path)

        if "error" in result:
            raise Exception(result["error"])

        return result.get("content", "")
    finally:
        # 删除临时文件
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)
